"""
Advanced Document Processor for RAG System
Supports multiple file types and web scraping
"""

import os
import re
import requests
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from urllib.parse import urlparse, urljoin
from urllib.robotparser import RobotFileParser

import fitz  # PyMuPDF for better PDF processing
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
import pandas as pd

# Simple Document class (matching the one in rag_system.py)
class Document:
    def __init__(self, page_content: str, metadata: Dict[str, Any] = None):
        self.page_content = page_content
        self.metadata = metadata or {}

    def __repr__(self):
        return f"Document(page_content='{self.page_content[:50]}...', metadata={self.metadata})"


logger = logging.getLogger(__name__)


class WebScraper:
    """Handles web scraping with respect for robots.txt"""

    def __init__(self, user_agent: str = "GenAI-Assistant/1.0"):
        self.user_agent = user_agent
        self.session = requests.Session()
        self.session.headers.update({"User-Agent": user_agent})

    def can_fetch(self, url: str) -> bool:
        """Check robots.txt before scraping"""
        try:
            parsed_url = urlparse(url)
            robots_url = f"{parsed_url.scheme}://{parsed_url.netloc}/robots.txt"

            rp = RobotFileParser()
            rp.set_url(robots_url)
            rp.read()

            return rp.can_fetch(self.user_agent, url)
        except Exception:
            # If we can't check robots.txt, be conservative
            return False

    def scrape_url(self, url: str, max_length: int = 50000) -> Optional[str]:
        """Scrape content from a URL"""
        try:
            if not self.can_fetch(url):
                logger.warning(f"Cannot fetch {url} - blocked by robots.txt")
                return None

            response = self.session.get(url, timeout=10)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()

            # Get text content
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)

            # Truncate if too long
            if len(text) > max_length:
                text = text[:max_length] + "..."

            return text

        except Exception as e:
            logger.error(f"Error scraping {url}: {e}")
            return None

    def scrape_multiple_urls(self, urls: List[str]) -> List[Document]:
        """Scrape multiple URLs"""
        documents = []
        for url in urls:
            content = self.scrape_url(url)
            if content:
                documents.append(Document(
                    page_content=content,
                    metadata={"source": url, "type": "web"}
                ))
        return documents


class AdvancedDocumentProcessor:
    """Advanced document processor with multiple format support"""

    def __init__(self):
        self.web_scraper = WebScraper()

    def load_text_file(self, file_path: str) -> List[Document]:
        """Load text from a file with encoding detection"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise UnicodeDecodeError("Could not decode file with any encoding")

            return [Document(page_content=text, metadata={"source": file_path, "type": "text"})]

        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            return []

    def load_pdf_file(self, file_path: str) -> List[Document]:
        """Load text from PDF using PyMuPDF (better than PyPDF2)"""
        try:
            doc = fitz.open(file_path)
            documents = []

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()

                # Extract metadata if available
                metadata = doc.metadata
                page_metadata = {
                    "source": file_path,
                    "type": "pdf",
                    "page": page_num + 1,
                    "total_pages": len(doc),
                    "title": metadata.get("title", ""),
                    "author": metadata.get("author", ""),
                    "subject": metadata.get("subject", "")
                }

                if text.strip():  # Only add non-empty pages
                    documents.append(Document(
                        page_content=text,
                        metadata=page_metadata
                    ))

            doc.close()
            return documents

        except Exception as e:
            logger.error(f"Error loading PDF file {file_path}: {e}")
            return []

    def load_docx_file(self, file_path: str) -> List[Document]:
        """Load text from Word document"""
        try:
            doc = DocxDocument(file_path)
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return [Document(page_content=text, metadata={"source": file_path, "type": "docx"})]

        except Exception as e:
            logger.error(f"Error loading DOCX file {file_path}: {e}")
            return []

    def load_csv_file(self, file_path: str) -> List[Document]:
        """Load data from CSV file"""
        try:
            df = pd.read_csv(file_path)

            # Convert DataFrame to text representation
            text = f"CSV File: {Path(file_path).name}\n\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n"
            text += f"Rows: {len(df)}\n\n"

            # Add first few rows as examples
            text += "Sample Data:\n"
            text += str(df.head().to_string()) + "\n\n"

            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text += "Summary Statistics:\n"
                text += str(df[numeric_cols].describe().to_string())

            return [Document(page_content=text, metadata={"source": file_path, "type": "csv"})]

        except Exception as e:
            logger.error(f"Error loading CSV file {file_path}: {e}")
            return []

    def load_markdown_file(self, file_path: str) -> List[Document]:
        """Load markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # Basic markdown cleaning (remove some formatting)
            text = re.sub(r'#+\s*', '', text)  # Remove headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
            text = re.sub(r'`([^`]+)`', r'\1', text)  # Remove inline code
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Remove links, keep text

            return [Document(page_content=text, metadata={"source": file_path, "type": "markdown"})]

        except Exception as e:
            logger.error(f"Error loading markdown file {file_path}: {e}")
            return []

    def load_from_url(self, url: str) -> List[Document]:
        """Load content from a URL"""
        content = self.web_scraper.scrape_url(url)
        if content:
            return [Document(page_content=content, metadata={"source": url, "type": "web"})]
        return []

    def load_from_multiple_urls(self, urls: List[str]) -> List[Document]:
        """Load content from multiple URLs"""
        return self.web_scraper.scrape_multiple_urls(urls)

    def load_documents_from_directory(self, directory: str, recursive: bool = True) -> List[Document]:
        """Load all supported documents from a directory"""
        documents = []
        directory_path = Path(directory)

        # Supported file extensions
        supported_extensions = {
            '.txt': self.load_text_file,
            '.pdf': self.load_pdf_file,
            '.docx': self.load_docx_file,
            '.doc': self.load_docx_file,  # .doc might work with python-docx
            '.csv': self.load_csv_file,
            '.md': self.load_markdown_file,
            '.markdown': self.load_markdown_file
        }

        # Use rglob for recursive search, glob for non-recursive
        pattern = "**/*" if recursive else "*"
        for file_path in directory_path.glob(pattern):
            if file_path.is_file():
                extension = file_path.suffix.lower()
                if extension in supported_extensions:
                    loader = supported_extensions[extension]
                    documents.extend(loader(str(file_path)))

        return documents

    def load_from_mixed_sources(self, sources: List[Union[str, Dict[str, Any]]]) -> List[Document]:
        """
        Load documents from mixed sources (files, directories, URLs)

        Args:
            sources: List of sources, each can be:
                - String path to file/directory
                - Dict with 'type' and 'path'/'url' keys
                - URL string (detected by http/https prefix)
        """
        documents = []

        for source in sources:
            if isinstance(source, str):
                # Detect if it's a URL
                if source.startswith(('http://', 'https://')):
                    documents.extend(self.load_from_url(source))
                elif os.path.isfile(source):
                    # Load single file
                    file_docs = self._load_single_file(source)
                    documents.extend(file_docs)
                elif os.path.isdir(source):
                    # Load directory
                    documents.extend(self.load_documents_from_directory(source))
                else:
                    logger.warning(f"Source not found: {source}")

            elif isinstance(source, dict):
                source_type = source.get('type', '').lower()
                if source_type == 'url':
                    url = source.get('url', source.get('path', ''))
                    documents.extend(self.load_from_url(url))
                elif source_type == 'file':
                    path = source.get('path', '')
                    if os.path.exists(path):
                        file_docs = self._load_single_file(path)
                        documents.extend(file_docs)
                elif source_type == 'directory':
                    path = source.get('path', '')
                    recursive = source.get('recursive', True)
                    documents.extend(self.load_documents_from_directory(path, recursive))

        return documents

    def _load_single_file(self, file_path: str) -> List[Document]:
        """Load a single file based on its extension"""
        path = Path(file_path)
        extension = path.suffix.lower()

        loaders = {
            '.txt': self.load_text_file,
            '.pdf': self.load_pdf_file,
            '.docx': self.load_docx_file,
            '.doc': self.load_docx_file,
            '.csv': self.load_csv_file,
            '.md': self.load_markdown_file,
            '.markdown': self.load_markdown_file
        }

        loader = loaders.get(extension)
        if loader:
            return loader(file_path)
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return []

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return ['txt', 'pdf', 'docx', 'doc', 'csv', 'md', 'markdown', 'web URLs']

    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """Get statistics about loaded documents"""
        stats = {
            'total_documents': len(documents),
            'total_characters': sum(len(doc.page_content) for doc in documents),
            'sources': {},
            'types': {}
        }

        for doc in documents:
            source = doc.metadata.get('source', 'unknown')
            doc_type = doc.metadata.get('type', 'unknown')

            stats['sources'][source] = stats['sources'].get(source, 0) + 1
            stats['types'][doc_type] = stats['types'].get(doc_type, 0) + 1

        return stats


# Convenience functions
def load_documents(sources: List[Union[str, Dict[str, Any]]]) -> List[Document]:
    """Convenience function to load documents from various sources"""
    processor = AdvancedDocumentProcessor()
    return processor.load_from_mixed_sources(sources)


def get_document_stats(documents: List[Document]) -> Dict[str, Any]:
    """Convenience function to get document statistics"""
    processor = AdvancedDocumentProcessor()
    return processor.get_document_stats(documents)